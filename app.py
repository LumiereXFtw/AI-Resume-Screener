from flask import Flask, request, jsonify
from flask_cors import CORS
import pdfplumber
import docx
import spacy
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import google.generativeai as genai
import os
import re
import logging
from datetime import datetime
import tempfile
from collections import Counter

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app)

# Load spaCy model with error handling
try:
    # Ensure you have downloaded the model: python -m spacy download en_core_web_sm
    nlp = spacy.load("en_core_web_sm")
    logger.info("spaCy model loaded successfully")
except OSError:
    logger.error("spaCy model 'en_core_web_sm' not found. Please install it with: python -m spacy download en_core_web_sm")
    nlp = None

# Configure Gemini API key from environment
# IMPORTANT: It's best practice to set this as an environment variable (e.g., export GEMINI_API_KEY="your_key")
# For local testing, you can uncomment the line below and hardcode, but do NOT do this in production.
GEMINI_API_KEY = ""
if not GEMINI_API_KEY:
    # Fallback for demonstration if environment variable is not set.
    # Replace "YOUR_GEMINI_API_KEY_HERE" with your actual key for testing.
    # REMEMBER TO REMOVE OR SECURELY MANAGE THIS IN PRODUCTION!
    GEMINI_API_KEY = ""
    logger.warning("GEMINI_API_KEY not found in environment variables. Using hardcoded default (for development only).")

if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)
    logger.info("Gemini API configured")
else:
    logger.error("Gemini API key is not configured. AI analysis will not be available.")


# In-memory storage for processed screenings (for demonstration, not production-ready)
# In a real application, you would use a database (SQL, NoSQL, etc.) like SQLAlchemy, MongoDB, etc.
processed_screenings_data = []

class ResumeAnalyzer:
    def __init__(self):
        self.skills_keywords = {
            'programming': ['python', 'java', 'javascript', 'c++', 'c#', 'ruby', 'php', 'go', 'rust', 'kotlin', 'swift'],
            'web_development': ['html', 'css', 'react', 'angular', 'vue', 'node.js', 'express', 'django', 'flask'],
            'data_science': ['pandas', 'numpy', 'matplotlib', 'scikit-learn', 'tensorflow', 'pytorch', 'sql', 'mongodb'],
            'cloud': ['aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform'],
            'project_management': ['agile', 'scrum', 'kanban', 'jira', 'confluence'],
            'databases': ['mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch'],
            'tools': ['git', 'jenkins', 'gitlab', 'github', 'vscode', 'intellij'],
            'soft_skills': ['leadership', 'communication', 'teamwork', 'problem-solving', 'analytical']
        }
        
    def extract_text(self, file):
        """Enhanced text extraction with better error handling"""
        try:
            # Create a temporary file to save the BytesIO object content
            # This is necessary because some libraries like pdfplumber and docx expect a file path or a file-like object that can be sought.
            # file.stream is a BytesIO object, but saving it to a temp file ensures compatibility.
            with tempfile.NamedTemporaryFile(delete=False, suffix=f"_{file.filename}") as temp_file:
                file.save(temp_file)
                temp_file_path = temp_file.name
            
            text_content = ""
            filename = file.filename.lower()
            
            try:
                if filename.endswith('.pdf'):
                    text_content = self._extract_from_pdf(temp_file_path)
                elif filename.endswith(('.docx', '.doc')):
                    text_content = self._extract_from_docx(temp_file_path)
                elif filename.endswith('.txt'):
                    with open(temp_file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        text_content = f.read()
                else:
                    # Generic fallback: try to read as plain text if extension is unknown
                    with open(temp_file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        text_content = f.read()
            finally:
                # Ensure the temporary file is deleted
                os.unlink(temp_file_path)
                
            return text_content
                
        except Exception as e:
            logger.error(f"Error extracting text from file '{file.filename}': {str(e)}")
            return ""
    
    def _extract_from_pdf(self, file_path):
        """Extract text from PDF with better handling"""
        try:
            text_content = []
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content.append(page_text.strip())
                        
                        # Also try to extract tables if present
                        tables = page.extract_tables()
                        for table in tables:
                            for row in table:
                                if row:
                                    text_content.append(' '.join([cell for cell in row if cell]))
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num} from PDF '{file_path}': {str(e)}")
                        continue
            
            return '\n'.join(text_content)
        except Exception as e:
            logger.error(f"Error opening or extracting PDF '{file_path}': {str(e)}")
            return ""
    
    def _extract_from_docx(self, file_path):
        """Extract text from DOCX with better handling"""
        try:
            doc = docx.Document(file_path)
            text_content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' '.join(row_text))
            
            return '\n'.join(text_content)
        except Exception as e:
            logger.error(f"Error opening or extracting DOCX '{file_path}': {str(e)}")
            return ""
    
    def preprocess_text(self, text):
        """Clean and preprocess text"""
        # Remove extra whitespace and normalize
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters but keep alphanumeric and common punctuation
        text = re.sub(r'[^\w\s\-\.\,\;\:\(\)\[\]]', ' ', text)
        return text.strip().lower()
    
    def extract_skills(self, text):
        """Extract skills from text using keyword matching"""
        text_lower = text.lower()
        found_skills = {}
        
        for category, skills in self.skills_keywords.items():
            found_skills[category] = []
            for skill in skills:
                # Use word boundaries to avoid partial matches (e.g., "java" in "javascript")
                if re.search(r'\b' + re.escape(skill) + r'\b', text_lower):
                    found_skills[category].append(skill)
        
        return found_skills
    
    def extract_experience_years(self, text):
        """Extract years of experience from text"""
        patterns = [
            r'(\d+)\+?\s*years?\s*of\s*experience',
            r'(\d+)\+?\s*years?\s*experience',
            r'experience\s*of\s*(\d+)\+?\s*years?',
            r'(\d+)\+?\s*yrs?\s*of\s*experience',
            r'(\d+)\s*y\.?\s*exp', # e.g., "5 y exp"
            r'(\d+)\s*year(?:s)?\s*[\-–—]\s*[\d.]+\s*year(?:s)?', # e.g., "2-5 years" for range
            r'total\s*experience\s*(\d+)',
            r'career\s*length\s*(\d+)'
        ]
        
        years = []
        for pattern in patterns:
            matches = re.findall(pattern, text.lower())
            for match in matches:
                # If the match is a tuple (e.g., from '(\+?\d{1,3})?\(?\d{3}\)?'), take the relevant part
                if isinstance(match, tuple):
                    try:
                        years.append(int(match[0]))
                    except ValueError:
                        pass # Ignore non-numeric parts
                else:
                    try:
                        years.append(int(match))
                    except ValueError:
                        pass # Ignore non-numeric parts

        if years:
            # Return the maximum year found, or a reasonable average/estimate
            return max(years)
        
        # Fallback: simple count of "years" or "experience" if specific patterns fail
        doc = nlp(text)
        for ent in doc.ents:
            if ent.label_ == "DATE" and ("year" in ent.text.lower() or "years" in ent.text.lower()):
                num_years = re.search(r'(\d+)', ent.text)
                if num_years:
                    years.append(int(num_years.group(1)))
        
        return max(years) if years else 0
    
    def extract_contact_info(self, text):
        """Extract contact information from resume"""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        # Improved phone number regex to be more flexible, matching common formats
        phone_pattern = r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}(?:\s*x\d+)?|\b\d{10}\b'
        # LinkedIn profile (basic detection)
        linkedin_pattern = r'(?:linkedin\.com/in/|linkedin\.com/mwlite/in/)([a-zA-Z0-9\-\_]+)'
        # GitHub profile (basic detection)
        github_pattern = r'(?:github\.com/)([a-zA-Z0-9\-\_]+)'
        
        emails = re.findall(email_pattern, text)
        phones = re.findall(phone_pattern, text)
        linkedins = re.findall(linkedin_pattern, text)
        githubs = re.findall(github_pattern, text)
        
        return {
            'emails': list(set(emails)), # Use set to remove duplicates
            'phones': list(set([re.sub(r'[^0-9+]', '', p) for p in phones])), # Clean phone numbers
            'linkedin_profiles': [f"linkedin.com/in/{p}" for p in list(set(linkedins))],
            'github_profiles': [f"github.com/{g}" for g in list(set(githubs))]
        }
    
    def extract_education(self, text):
        """Extract education information"""
        education_keywords = ['bachelor', 'master', 'phd', 'degree', 'university', 'college', 'institute', 'graduated', 'major', 'minor', 'diploma']
        education_info = []
        
        # Split text into lines or sentences to process chunks
        # Using spaCy for sentence segmentation can be more robust
        if nlp:
            doc = nlp(text)
            for sent in doc.sents:
                sent_lower = sent.text.lower()
                if any(keyword in sent_lower for keyword in education_keywords):
                    education_info.append(sent.text.strip())
        else: # Fallback to line splitting if spaCy is not loaded
            lines = text.split('\n')
            for line in lines:
                line_lower = line.lower()
                if any(keyword in line_lower for keyword in education_keywords):
                    education_info.append(line.strip())
        
        return education_info[:10]  # Return top 10 education-related lines/sentences
    
    def calculate_advanced_similarity(self, resume_text, job_description):
        """Calculate multiple similarity metrics"""
        try:
            # TF-IDF Similarity
            vectorizer = TfidfVectorizer(
                stop_words='english',
                ngram_range=(1, 2),
                max_features=2000 # Increased max features for potentially richer representation
            )
            tfidf_matrix = vectorizer.fit_transform([job_description, resume_text])
            tfidf_similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
            
            # spaCy Semantic Similarity
            spacy_similarity = 0
            if nlp:
                try:
                    # Limiting text length for spaCy for performance, though 1M is already quite large
                    resume_doc = nlp(resume_text[:1000000])
                    jd_doc = nlp(job_description[:1000000])
                    spacy_similarity = resume_doc.similarity(jd_doc)
                except Exception as e:
                    logger.warning(f"spaCy similarity calculation failed (might be due to model limitations or text length): {str(e)}")
                    spacy_similarity = 0
            
            # Keyword Matching
            jd_keywords = set(self.preprocess_text(job_description).split())
            resume_words = set(self.preprocess_text(resume_text).split())
            
            # Filter out common words (expanded list)
            common_words = {'the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'from', 'up', 'about', 'into', 'through', 'during', 'before', 'after', 'above', 'below', 'between', 'among', 'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might', 'must', 'can', 'shall', 'a', 'an', 'as', 'i', 'he', 'she', 'it', 'we', 'you', 'they', 'my', 'your', 'his', 'her', 'its', 'our', 'their', 'this', 'that', 'these', 'those', 'here', 'there', 'when', 'where', 'why', 'how', 'all', 'any', 'both', 'each', 'few', 'more', 'most', 'other', 'some', 'such', 'no', 'nor', 'not', 'only', 'own', 'same', 'so', 'than', 'too', 'very', 's', 't', 'can', 'will', 'just', 'don', 'should', 'now'}
            
            jd_keywords = jd_keywords - common_words
            resume_words = resume_words - common_words
            
            matched_keywords = list(jd_keywords & resume_words)
            keyword_match_ratio = len(matched_keywords) / len(jd_keywords) if jd_keywords else 0
            
            # Skill-specific matching
            resume_skills = self.extract_skills(resume_text)
            jd_skills = self.extract_skills(job_description)
            
            skill_matches = []
            for category in resume_skills:
                common_skills = set(resume_skills[category]) & set(jd_skills.get(category, []))
                skill_matches.extend(list(common_skills))
            
            return {
                'tfidf_similarity': tfidf_similarity,
                'spacy_similarity': spacy_similarity,
                'keyword_match_ratio': keyword_match_ratio,
                'matched_keywords': matched_keywords,
                'skill_matches': skill_matches,
                'resume_skills': resume_skills,
                'jd_skills': jd_skills
            }
            
        except Exception as e:
            logger.error(f"Error calculating similarity: {str(e)}")
            return {
                'tfidf_similarity': 0,
                'spacy_similarity': 0,
                'keyword_match_ratio': 0,
                'matched_keywords': [],
                'skill_matches': [],
                'resume_skills': {},
                'jd_skills': {}
            }
    
    def analyze_with_gemini(self, resume_text, job_description):
        """Performs comprehensive Gemini analysis, including overall score and recommendation."""
        if not GEMINI_API_KEY:
            return {
                "ai_analysis_text": "Gemini API key not configured. Skipping AI analysis.",
                "overall_score_gemini": 0,
                "overall_recommendation_gemini": "Not Available"
            }
        
        try:
            # Directly specify the stable model to use
            model_to_use = "gemini-2.0-flash" # Using gemini-pro as a stable, widely available model

            logger.info(f"Attempting to use Gemini model: {model_to_use}")
            model = genai.GenerativeModel(model_to_use)

            prompt = f"""
You are an expert HR recruiter and career coach. Analyze this resume against the job description and provide comprehensive feedback.

First, provide a numerical Overall Match Score (0-100%) and an Overall Recommendation string.
Then, provide detailed analysis covering the points below.

Overall Match Score (0-100%):
Overall Recommendation:

RESUME:
{resume_text[:3000]}

JOB DESCRIPTION:
{job_description[:2000]}

Detailed Analysis (sections and bullet points):
1. Key strengths that directly match the role requirements.
2. Skills alignment: List matched skills and identify any significant skill gaps.
3. Experience relevance: How well does the candidate's experience align with the job responsibilities?
4. Areas for improvement: Suggestions for the candidate to improve their resume or skills for future applications.
5. Red flags or concerns: Any potential issues (e.g., significant gaps, lack of relevant experience).
6. Recommended interview focus areas: Specific topics or skills to probe during an interview.
7. Salary range estimate based on experience and role (e.g., $X,000 - $Y,000 annually).
"""
            
            response = model.generate_content(prompt)
            ai_analysis_text = response.text.strip()

            overall_score_gemini = 0
            overall_recommendation_gemini = "Undetermined"

            # Parse score and recommendation from the beginning of Gemini's response
            score_match = re.search(r'Overall Match Score \(0-100%\):\s*(\d+)', ai_analysis_text)
            if score_match:
                try:
                    overall_score_gemini = int(score_match.group(1))
                    # Remove the matched line to clean up the main analysis text
                    ai_analysis_text = re.sub(r'Overall Match Score \(0-100%\):\s*\d+\s*', '', ai_analysis_text, 1).strip()
                except ValueError:
                    logger.warning("Could not parse Gemini's overall score.")
            
            recommendation_match = re.search(r'Overall Recommendation:\s*(.+)', ai_analysis_text.split('\n')[0] if ai_analysis_text else '', re.IGNORECASE)
            if recommendation_match:
                overall_recommendation_gemini = recommendation_match.group(1).strip()
                # Remove the matched line to clean up the main analysis text
                ai_analysis_text = re.sub(r'Overall Recommendation:\s*.+\s*', '', ai_analysis_text, 1).strip()
            
            return {
                "ai_analysis_text": ai_analysis_text,
                "overall_score_gemini": overall_score_gemini,
                "overall_recommendation_gemini": overall_recommendation_gemini
            }
            
        except Exception as e:
            logger.error(f"Gemini API call failed for '{model_to_use}': {str(e)}")
            return {
                "ai_analysis_text": f"Error generating AI analysis: Failed to use '{model_to_use}'. Reason: {str(e)}",
                "overall_score_gemini": 0,
                "overall_recommendation_gemini": "Error during AI analysis"
            }
    
    def get_recommendation(self, composite_score):
        """This function will no longer be the primary source of recommendation if Gemini provides one."""
        # This function might be deprecated or used as a fallback if Gemini parsing fails.
        if composite_score >= 80:
            return "Strong Match - Highly Recommended"
        elif composite_score >= 65:
            return "Good Match - Recommended"
        elif composite_score >= 50:
            return "Moderate Match - Consider with reservations"
        elif composite_score >= 35:
            return "Weak Match - Not recommended"
        else:
            return "Poor Match - Reject"
    
    def _calculate_skills_gap(self, resume_skills, job_skills):
        """Calculate skills gap between resume and job requirements"""
        gap = {}
        for category in job_skills:
            required_skills = set(job_skills[category])
            candidate_skills = set(resume_skills.get(category, []))
            missing_skills = required_skills - candidate_skills
            if missing_skills: # Only add categories with actual missing skills
                gap[category] = list(missing_skills)
        return gap

# Initialize analyzer
analyzer = ResumeAnalyzer()

@app.route('/api/screen', methods=['POST'])
def screen_resume():
    """Enhanced resume screening endpoint"""
    try:
        # Validate request
        if 'resume' not in request.files:
            return jsonify({"error": "No resume file provided"}), 400
        
        if 'job_description' not in request.form:
            return jsonify({"error": "No job description provided"}), 400
        
        resume_file = request.files['resume']
        job_description = request.form['job_description'].strip()
        
        if not resume_file.filename:
            return jsonify({"error": "No file selected"}), 400
        
        if not job_description:
            return jsonify({"error": "Job description cannot be empty"}), 400
        
        logger.info(f"Processing resume: {resume_file.filename}")
        
        # Extract text from resume
        resume_text = analyzer.extract_text(resume_file)
        if not resume_text:
            return jsonify({"error": f"Could not extract text from resume '{resume_file.filename}'. Please ensure it's a valid PDF, DOCX, or TXT file."}), 400
        
        logger.info(f"Extracted {len(resume_text)} characters from resume '{resume_file.filename}'")
        
        # Perform advanced analysis (Python-based for raw data/keyword matching)
        similarity_metrics = analyzer.calculate_advanced_similarity(resume_text, job_description)
        
        # Extract additional insights (Python-based)
        experience_years = analyzer.extract_experience_years(resume_text)
        contact_info = analyzer.extract_contact_info(resume_text)
        education_info = analyzer.extract_education(resume_text)
        
        # Get Gemini analysis - this is now the primary source of judgment
        gemini_result = analyzer.analyze_with_gemini(resume_text, job_description)
        gemini_analysis = gemini_result["ai_analysis_text"]
        composite_score_gemini = gemini_result["overall_score_gemini"]
        recommendation_gemini = gemini_result["overall_recommendation_gemini"]

        # Calculate Python-based scores for informational purposes (no longer used for primary recommendation)
        tfidf_score = round(similarity_metrics['tfidf_similarity'] * 100, 2)
        spacy_score = round(similarity_metrics['spacy_similarity'] * 100, 2)
        keyword_score = round(similarity_metrics['keyword_match_ratio'] * 100, 2)
        
        # Prepare response
        result = {
            "success": True,
            "filename": resume_file.filename,
            "job_description_summary": job_description[:100] + ("..." if len(job_description) > 100 else ""), # Short summary
            "processed_at": datetime.now().isoformat(),
            "scores": {
                "tfidf_score": tfidf_score,
                "spacy_score": spacy_score,
                "keyword_score": keyword_score,
                "composite_score": composite_score_gemini # Composite score now comes from Gemini
            },
            "matching": {
                "matched_skills": similarity_metrics['skill_matches'],
                "matched_keywords": similarity_metrics['matched_keywords'][:30],  # Limit to top 30
                "total_matched_keywords": len(similarity_metrics['matched_keywords'])
            },
            "skills_analysis": {
                "resume_found_skills": {k: v for k, v in similarity_metrics['resume_skills'].items() if v}, # Only show categories with found skills
                "job_required_skills": {k: v for k, v in similarity_metrics['jd_skills'].items() if v}, # Only show categories with required skills
                "skills_gap": analyzer._calculate_skills_gap(
                    similarity_metrics['resume_skills'],
                    similarity_metrics['jd_skills']
                )
            },
            "candidate_info": {
                "experience_years": experience_years,
                "contact_info": contact_info,
                "education": education_info,
                "resume_length_chars": len(resume_text),
                "resume_length_words": len(resume_text.split())
            },
            "recommendation": {
                "overall": recommendation_gemini, # Overall recommendation now comes from Gemini
                "score_breakdown": { # Breakdown based on Gemini's score
                    "gemini_based": "High" if composite_score_gemini >= 70 else "Medium" if composite_score_gemini >= 40 else "Low"
                }
            },
            "ai_analysis": gemini_analysis, # Full detailed AI analysis
            "metadata": {
                "processing_time": "N/A",  # Implement actual timing if needed
                "model_versions": {
                    "spacy": "available" if nlp else "not_available",
                    "gemini": "available" if GEMINI_API_KEY else "not_available" # This checks if API key is set
                }
            }
        }
        
        # Store the result for retrieval by /api/screenings (for demonstration)
        processed_screenings_data.append(result)

        logger.info(f"Analysis complete for {resume_file.filename} - Gemini Score: {composite_score_gemini}")
        return jsonify(result)
        
    except Exception as e:
        logger.error(f"Error in resume screening: {str(e)}")
        return jsonify({
            "success": False,
            "error": f"Internal server error during screening: {str(e)}"
        }), 500

@app.route('/api/batch_screen', methods=['POST'])
def batch_screen_resumes():
    """Batch process multiple resumes"""
    try:
        if 'resumes' not in request.files:
            return jsonify({"error": "No resume files provided"}), 400
        
        if 'job_description' not in request.form:
            return jsonify({"error": "No job description provided"}), 400
        
        resume_files = request.files.getlist('resumes')
        job_description = request.form['job_description'].strip()
        
        if not resume_files:
            return jsonify({"error": "No files selected"}), 400
        
        if not job_description:
            return jsonify({"error": "Job description cannot be empty"}), 400

        logger.info(f"Starting batch processing for {len(resume_files)} resumes.")
        
        results = []
        for resume_file in resume_files:
            file_result = {
                "filename": resume_file.filename,
                "processed_at": datetime.now().isoformat()
            }
            try:
                resume_text = analyzer.extract_text(resume_file)
                if resume_text:
                    # Perform Python-based analysis for basic metrics
                    similarity_metrics = analyzer.calculate_advanced_similarity(resume_text, job_description)

                    # Get Gemini analysis for judgment
                    gemini_result = analyzer.analyze_with_gemini(resume_text, job_description)
                    composite_score_gemini = gemini_result["overall_score_gemini"]
                    recommendation_gemini = gemini_result["overall_recommendation_gemini"]
                    
                    file_result.update({
                        "composite_score": composite_score_gemini, # Use Gemini's score
                        "recommendation": recommendation_gemini, # Use Gemini's recommendation
                        "matched_skills": similarity_metrics['skill_matches'], # Python-based
                        "experience_years": analyzer.extract_experience_years(resume_text), # Python-based
                        "success": True,
                        "ai_analysis_summary": gemini_result["ai_analysis_text"][:200] + "..." if len(gemini_result["ai_analysis_text"]) > 200 else gemini_result["ai_analysis_text"]
                    })
                    # Store a more concise version for batch/list views
                    processed_screenings_data.append({
                        "filename": resume_file.filename,
                        "job_description_summary": job_description[:100] + ("..." if len(job_description) > 100 else ""),
                        "processed_at": file_result['processed_at'],
                        "scores": { "composite_score": composite_score_gemini },
                        "recommendation": { "overall": recommendation_gemini },
                        "matching": { "matched_skills": file_result['matched_skills'] },
                        "candidate_info": { "experience_years": file_result['experience_years'] },
                        "success": True
                    })
                else:
                    file_result.update({
                        "error": "Could not extract text",
                        "success": False
                    })
            except Exception as e:
                logger.error(f"Error processing {resume_file.filename} in batch: {str(e)}")
                file_result.update({
                    "error": str(e),
                    "success": False
                })
            results.append(file_result)
        
        # Sort successful results by composite score (descending)
        successful_results = [r for r in results if r.get('success')]
        successful_results.sort(key=lambda x: x.get('composite_score', 0), reverse=True)
        
        logger.info(f"Batch processing complete. Successful: {len(successful_results)}, Failed: {len(results) - len(successful_results)}")

        return jsonify({
            "success": True,
            "total_processed": len(results),
            "successful": len(successful_results),
            "failed": len(results) - len(successful_results),
            "results": results, # Detailed results for each file in the batch
            "ranked_candidates": successful_results # Ranked list of successful candidates
        })
        
    except Exception as e:
        logger.error(f"Top-level error in batch screening: {str(e)}")
        return jsonify({
            "success": False,
            "error": f"Batch processing error: {str(e)}"
        }), 500

@app.route('/api/screenings', methods=['GET'])
def get_screenings():
    """Returns a list of all previously processed screenings."""
    page = request.args.get('page', 1, type=int)
    limit = request.args.get('limit', 10, type=int)
    sort_by = request.args.get('sortBy', 'processed_at')
    order = request.args.get('order', 'desc')

    try:
        sorted_data = sorted(
            processed_screenings_data,
            key=lambda x: x.get(sort_by, x.get('scores',{}).get('composite_score', 0) if sort_by == 'composite_score' else ''),
            reverse=(order == 'desc')
        )
    except TypeError:
         logger.warning(f"Sorting error with sortBy={sort_by}. Falling back to processed_at.")
         sorted_data = sorted(processed_screenings_data, key=lambda x: x.get('processed_at', ''), reverse=(order == 'desc'))


    start_index = (page - 1) * limit
    end_index = start_index + limit
    paginated_data = sorted_data[start_index:end_index]

    return jsonify({
        "success": True,
        "total_screenings": len(processed_screenings_data),
        "page": page,
        "limit": limit,
        "screenings": paginated_data,
        "message": "Note: Data is stored in-memory and will be lost on server restart. Implement a database for persistence."
    }), 200

@app.route('/api/analytics', methods=['GET'])
def get_analytics():
    """Returns aggregated analytics based on processed screenings."""
    total_resumes_processed = len(processed_screenings_data)
    
    # Calculate average composite score (now based on Gemini's score)
    valid_scores = [s['scores']['composite_score'] for s in processed_screenings_data
                    if 'scores' in s and 'composite_score' in s['scores'] and s.get('success', False)]
    average_composite_score = round(sum(valid_scores) / len(valid_scores), 2) if valid_scores else 0

    # Count recommendations (now based on Gemini's recommendation)
    recommendation_counts = Counter(s['recommendation']['overall'] for s in processed_screenings_data
                                    if 'recommendation' in s and 'overall' in s['recommendation'] and s.get('success', False))

    # Aggregate matched skills (top N) - still based on Python extraction
    all_matched_skills = []
    for s in processed_screenings_data:
        if s.get('success', False) and 'matching' in s and 'matched_skills' in s['matching']:
            all_matched_skills.extend(s['matching']['matched_skills'])
    top_matched_skills = Counter(all_matched_skills).most_common(10)

    # Aggregate skills gap (top N missing skills) - still based on Python extraction
    all_missing_skills = []
    for s in processed_screenings_data:
        if s.get('success', False) and 'skills_analysis' in s and 'skills_gap' in s['skills_analysis']:
            for category, skills in s['skills_analysis']['skills_gap'].items():
                all_missing_skills.extend(skills)
    top_missing_skills = Counter(all_missing_skills).most_common(10)

    # Experience distribution - still based on Python extraction
    experience_years_list = [s['candidate_info']['experience_years'] for s in processed_screenings_data
                             if s.get('success', False) and 'candidate_info' in s and 'experience_years' in s['candidate_info']]
    experience_distribution = Counter(experience_years_list).most_common()
    experience_distribution_dict = {str(k): v for k, v in sorted(experience_distribution)}


    return jsonify({
        "success": True,
        "total_resumes_processed": total_resumes_processed,
        "average_composite_score": average_composite_score,
        "recommendation_breakdown": dict(recommendation_counts),
        "top_matched_skills": top_matched_skills,
        "top_missing_skills": top_missing_skills,
        "experience_distribution": experience_distribution_dict,
        "message": "Note: Analytics are based on in-memory data and will reset on server restart. Implement a database for comprehensive analytics."
    }), 200

@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "services": {
            "spacy": nlp is not None,
            "gemini": GEMINI_API_KEY is not None
        }
    })

@app.route('/api/skills', methods=['GET'])
def get_available_skills():
    """Get list of available skill categories and keywords"""
    return jsonify({
        "skill_categories": analyzer.skills_keywords
    })

# Error handlers
@app.errorhandler(400)
def bad_request(error):
    logger.error(f"Bad request: {error.description}")
    return jsonify({"error": error.description or "Bad request"}), 400

@app.errorhandler(500)
def internal_error(error):
    logger.exception("Internal server error caught by error handler:")
    return jsonify({"error": "Internal server error. Please check server logs for details."}), 500

if __name__ == '__main__':
    # Set up the application
    port = int(os.environ.get('FLASK_PORT', 5001))
    debug = os.environ.get('FLASK_DEBUG', 'False').lower() == 'true'

    logger.info(f"Starting Resume Screening API (Python) on port {port}")
    logger.info(f"Debug mode: {debug}")
    logger.info(f"spaCy available: {nlp is not None}")
    logger.info(f"Gemini API available: {GEMINI_API_KEY is not None}")

    app.run(
        host='0.0.0.0',
        port=port,
        debug=debug
    )